"""
RAG Adaptativo para GMV Sistema
Integração com Llama31 via Ollama + LangChain
Versão: 1.2 - Final com langchain-ollama
"""

import os
import re
import json
import logging
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Any
from dataclasses import dataclass
from enum import Enum

# Core dependencies
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

# LangChain (versão mais atual)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate

rag_status = False

# LangChain Ollama (nova versão)
try:
    from langchain_ollama import OllamaLLM, OllamaEmbeddings
    OLLAMA_NEW = True
except ImportError:
    # Fallback para versão anterior
    try:
        from langchain_community.llms import Ollama as OllamaLLM
        from langchain_community.embeddings import OllamaEmbeddings
        OLLAMA_NEW = False
    except ImportError:
        from langchain.llms import Ollama as OllamaLLM
        from langchain.embeddings import OllamaEmbeddings
        OLLAMA_NEW = False

# Vector Store
try:
    from langchain_community.vectorstores import FAISS
except ImportError:
    from langchain.vectorstores import FAISS

# Document processing
import pypdf
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

# Local utilities
try:
    from utils.suspeicao import encontrar_suspeitos
except ImportError:
    def encontrar_suspeitos(text, file_path):
        """Fallback function for suspeitos detection"""
        return []

logger = logging.getLogger(__name__)

class QueryType(Enum):
    """Tipos de consulta para estratégia adaptativa"""
    FACTUAL = "factual"      # Busca fatos específicos
    ANALYTICAL = "analytical" # Análise abrangente 
    OPINION = "opinion"      # Múltiplas perspectivas
    CONTEXTUAL = "contextual" # Específico do usuário

@dataclass
class RAGConfig:
    """Configuração do sistema RAG"""
    model_name: str = "llama3.1:8b"
    embedding_model: str = "nomic-embed-text"
    chunk_size: int = 500
    chunk_overlap: int = 50
    top_k: int = 4
    temperature: float = 0.1
    max_tokens: int = 1024
    
class DocumentProcessor:
    """Processador de documentos para extração de texto"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extrai texto de arquivo PDF"""
        try:
            with open(file_path, 'rb') as file:
                reader = pypdf.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Erro ao extrair PDF {file_path}: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extrai texto de arquivo DOCX"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Erro ao extrair DOCX {file_path}: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extrai texto de arquivo TXT/MD"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Erro ao extrair TXT {file_path}: {e}")
            return ""
    
    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """Extrai texto baseado na extensão do arquivo"""
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return cls.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return cls.extract_text_from_docx(file_path)
        elif ext in ['.txt', '.md']:
            return cls.extract_text_from_txt(file_path)
        else:
            logger.warning(f"Tipo de arquivo não suportado: {ext}")
            return ""

class QueryClassifier:
    """Classificador de tipos de consulta para estratégia adaptativa"""
    
    # Palavras-chave para cada tipo de consulta
    FACTUAL_KEYWORDS = [
        'quando', 'onde', 'quem', 'qual', 'quanto', 'data', 'número',
        'nome', 'valor', 'endereço', 'telefone', 'cpf', 'cnpj'
    ]
    
    ANALYTICAL_KEYWORDS = [
        'análise', 'comparar', 'avaliar', 'examinar', 'investigar',
        'explicar', 'como', 'por que', 'causas', 'consequências'
    ]
    
    OPINION_KEYWORDS = [
        'opinião', 'perspectiva', 'visão', 'ponto de vista', 'considera',
        'acredita', 'pensa', 'diferentes', 'alternativas', 'possibilidades'
    ]
    
    CONTEXTUAL_KEYWORDS = [
        'meu caso', 'minha situação', 'para mim', 'no meu contexto',
        'considerando', 'levando em conta', 'específico', 'particular'
    ]
    
    @classmethod
    def classify_query(cls, query: str) -> QueryType:
        """Classifica o tipo da consulta"""
        query_lower = query.lower()
        
        # Conta ocorrências de palavras-chave
        scores = {
            QueryType.FACTUAL: sum(1 for word in cls.FACTUAL_KEYWORDS if word in query_lower),
            QueryType.ANALYTICAL: sum(1 for word in cls.ANALYTICAL_KEYWORDS if word in query_lower),
            QueryType.OPINION: sum(1 for word in cls.OPINION_KEYWORDS if word in query_lower),
            QueryType.CONTEXTUAL: sum(1 for word in cls.CONTEXTUAL_KEYWORDS if word in query_lower)
        }
        
        # Retorna o tipo com maior pontuação
        max_type = max(scores.items(), key=lambda x: x[1])
        
        # Se não há palavras-chave claras, usa factual como padrão
        return max_type[0] if max_type[1] > 0 else QueryType.FACTUAL

class AdaptiveRAG:
    """Sistema RAG Adaptativo principal"""
    
    def __init__(self, config: RAGConfig = None):
        self.config = config or RAGConfig()
        self.llm = None
        self.embeddings = None
        self.vector_store = None
        self.documents = []
        self.is_initialized = False
        
        # Pasta de dados do projeto - usa variável de ambiente se disponível
        self.data_path = os.getenv("PASTA_DESTINO", os.getenv("DADOS_ANONIMOS", "./documentos"))
        
    def initialize(self) -> bool:
        """Inicializa os modelos Ollama"""
        try:
            logger.info(" Inicializando RAG Adaptativo...")
            
            # Inicializa LLM (usando método mais atual)
            if OLLAMA_NEW:
                logger.info(" Usando langchain-ollama (versão atual)")
                self.llm = OllamaLLM(
                    model=self.config.model_name,
                    temperature=self.config.temperature,
                    num_predict=self.config.max_tokens
                )
            else:
                logger.info("⚠️ Usando langchain-community (versão anterior)")
                self.llm = OllamaLLM(
                    model=self.config.model_name,
                    temperature=self.config.temperature,
                    num_predict=self.config.max_tokens
                )
            
            # Inicializa embeddings
            self.embeddings = OllamaEmbeddings(
                model=self.config.embedding_model
            )
            
            # Testa conexão (usando invoke em vez de __call__)
            try:
                if hasattr(self.llm, 'invoke'):
                    test_response = self.llm.invoke("Teste de conexão. Responda apenas 'OK'.")
                else:
                    test_response = self.llm("Teste de conexão. Responda apenas 'OK'.")
                    
                logger.info(f" Ollama conectado: {test_response[:20]}...")
            except Exception as e:
                logger.warning(f"⚠️ Teste de conexão falhou: {e}")
                # Continua mesmo com falha no teste
            
            self.is_initialized = True
            return True
            
        except Exception as e:
            logger.error(f"❌ Erro ao inicializar Ollama: {e}")
            logger.error("💡 Verifique se Ollama está rodando: ollama serve")
            return False
    
    def load_documents_from_directory(self, directory: str = None) -> int:
        """Carrega documentos de um diretório"""
        if not self.is_initialized:
            logger.error("❌ RAG não inicializado!")
            return 0
            
        directory = directory or self.data_path
        if not os.path.exists(directory):
            logger.warning(f"⚠️ Diretório não existe: {directory}")
            return 0
        
        documents = []
        processed_files = 0
        
        logger.info(f" Carregando documentos de: {directory}")
        
        # Processa todos os arquivos suportados
        for root, dirs, files in os.walk(directory):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx', '.txt', '.md')):
                    file_path = os.path.join(root, file)
                    
                    try:
                        text = DocumentProcessor.extract_text(file_path)
                        if text.strip():
                            # Cria documento LangChain
                            doc = Document(
                                page_content=text,
                                metadata={
                                    "source": file_path,
                                    "filename": file,
                                    "size": len(text)
                                }
                            )
                            documents.append(doc)
                            processed_files += 1
                            logger.info(f" Processado: {file}")
                        else:
                            logger.warning(f"⚠️ Arquivo vazio: {file}")
                            
                    except Exception as e:
                        logger.error(f"❌ Erro ao processar {file}: {e}")
        
        if documents:
            self.documents = documents
            self._create_vector_store()
        else:
            logger.warning("⚠️ Nenhum documento foi carregado!")
            
        logger.info(f"📊 Total processado: {processed_files} arquivos")
        return processed_files
    
    def _create_vector_store(self):
        """Cria o vector store com chunking"""
        if not self.documents:
            logger.warning("⚠️ Nenhum documento carregado!")
            return
            
        logger.info(" Criando chunks e embeddings...")
        
        try:
            # Splitter para chunking
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.config.chunk_size,
                chunk_overlap=self.config.chunk_overlap,
                separators=["\n\n", "\n", ". ", " ", ""]
            )
            
            # Cria chunks
            chunks = text_splitter.split_documents(self.documents)
            logger.info(f"Criados {len(chunks)} chunks")
            global rag_status
            rag_status = True
            
            # Cria vector store
            self.vector_store = FAISS.from_documents(chunks, self.embeddings)
            logger.info(" Vector store criado!")
            
        except Exception as e:
            logger.error(f"❌ Erro ao criar vector store: {e}")
            raise
    
    def _get_retrieval_strategy(self, query_type: QueryType) -> Dict[str, Any]:
        """Define estratégia de retrieval baseada no tipo de consulta"""
        strategies = {
            QueryType.FACTUAL: {
                "k": 3,  # Menos documentos, mais precisos
                "search_type": "similarity",
                "score_threshold": 0.7
            },
            QueryType.ANALYTICAL: {
                "k": 6,  # Mais documentos para análise abrangente
                "search_type": "similarity",
                "score_threshold": 0.6
            },
            QueryType.OPINION: {
                "k": 5,  # Documentos variados
                "search_type": "mmr",  # Maximum Marginal Relevance
                "score_threshold": 0.5
            },
            QueryType.CONTEXTUAL: {
                "k": 4,  # Balanceado
                "search_type": "similarity",
                "score_threshold": 0.6
            }
        }
        return strategies.get(query_type, strategies[QueryType.FACTUAL])
    
    def _get_prompt_template(self, query_type: QueryType) -> str:
        """Define template de prompt baseado no tipo de consulta"""
        templates = {
            QueryType.FACTUAL: """
Com base nos documentos fornecidos, responda à pergunta de forma factual e precisa.
Se a informação não estiver disponível, diga claramente.

Documentos:
{context}

Pergunta: {question}

Resposta factual:""",
            
            QueryType.ANALYTICAL: """
Analise os documentos fornecidos e forneça uma resposta abrangente e analítica.
Explore diferentes aspectos e conexões entre as informações.

Documentos:
{context}

Pergunta: {question}

Análise detalhada:""",
            
            QueryType.OPINION: """
Com base nos documentos, apresente diferentes perspectivas e pontos de vista sobre a questão.
Considere múltiplas interpretações quando aplicável.

Documentos:
{context}

Pergunta: {question}

Múltiplas perspectivas:""",
            
            QueryType.CONTEXTUAL: """
Considerando o contexto específico mencionado, forneça uma resposta personalizada
baseada nos documentos disponíveis.

Documentos:
{context}

Pergunta: {question}

Resposta contextualizada:"""
        }
        return templates.get(query_type, templates[QueryType.FACTUAL])
    
    def query(self, question: str, context: str = None) -> Dict[str, Any]:
        """Executa consulta RAG adaptativa"""
        if not self.is_initialized or not self.vector_store:
            return {
                "error": "RAG não inicializado ou sem documentos carregados"
            }
        
        try:
            # 1. Classifica o tipo da consulta
            query_type = QueryClassifier.classify_query(question)
            logger.info(f"🎯 Consulta classificada como: {query_type.value}")
            
            # 2. Define estratégia de retrieval
            strategy = self._get_retrieval_strategy(query_type)
            
            # 3. Busca documentos relevantes
            retriever = self.vector_store.as_retriever(
                search_kwargs={"k": strategy["k"]}
            )
            
            # 4. Cria prompt adaptativo
            prompt_template = self._get_prompt_template(query_type)
            prompt = PromptTemplate(
                template=prompt_template,
                input_variables=["context", "question"]
            )
            
            # 5. Executa chain RAG
            qa_chain = RetrievalQA.from_chain_type(
                llm=self.llm,
                chain_type="stuff",
                retriever=retriever,
                chain_type_kwargs={"prompt": prompt},
                return_source_documents=True
            )
            
            result = qa_chain({"query": question})
            
            # 6. Processa suspeitos se relevante
            suspeitos = []
            if "suspeito" in question.lower() or "suspeita" in question.lower():
                suspeitos = encontrar_suspeitos(result["result"], "")
            
            return {
                "question": question,
                "query_type": query_type.value,
                "answer": result["result"],
                "source_documents": [
                    {
                        "content": doc.page_content[:200] + "...",
                        "source": doc.metadata.get("source", ""),
                        "filename": doc.metadata.get("filename", "")
                    }
                    for doc in result["source_documents"]
                ],
                "suspeitos": suspeitos,
                "strategy_used": strategy,
                "processing_time": 0.0  # Pode ser implementado se necessário
            }
            
        except Exception as e:
            logger.error(f"❌ Erro na consulta RAG: {e}")
            return {"error": str(e)}
    
    def add_document_text(self, text: str, metadata: Dict = None) -> bool:
        """Adiciona texto como documento"""
        if not self.is_initialized:
            return False
            
        try:
            doc = Document(
                page_content=text,
                metadata=metadata or {}
            )
            
            self.documents.append(doc)
            self._create_vector_store()
            return True
            
        except Exception as e:
            logger.error(f"❌ Erro ao adicionar documento: {e}")
            return False
    
    def get_stats(self) -> Dict[str, Any]:
        """Retorna estatísticas do sistema"""
        stats = {
            "initialized": self.is_initialized,
            "documents_count": len(self.documents),
            "has_vector_store": self.vector_store is not None,
            "vector_store_size": 0,
            "ollama_version": "langchain-ollama" if OLLAMA_NEW else "langchain-community",
            "config": {
                "model": self.config.model_name,
                "embedding_model": self.config.embedding_model,
                "chunk_size": self.config.chunk_size,
                "data_path": self.data_path
            }
        }
        
        # Calcula tamanho do vector store se disponível
        if self.vector_store:
            try:
                # FAISS não tem método direto para contar, estimamos pelos documentos
                stats["vector_store_size"] = len(self.documents)
            except Exception:
                stats["vector_store_size"] = 0
        
        return stats

# Instância global para uso no Flask
rag_system = AdaptiveRAG()

def init_rag_system() -> bool:
    """Inicializa o sistema RAG global"""
    return rag_system.initialize()

def load_data_directory() -> int:
    """Carrega dados do diretório configurado"""
    return rag_system.load_documents_from_directory()